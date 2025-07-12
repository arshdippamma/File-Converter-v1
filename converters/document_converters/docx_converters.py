import os
from docx import Document
from docx2pdf import convert

def docx_to_pdf(input_path):
    output_path = os.path.splitext(input_path)[0] + ".pdf"
    convert(input_path, output_path)
    return output_path

def docx_to_txt(input_path):
    output_path = os.path.splitext(input_path)[0] + ".txt"

    lines = []

    doc = Document(input_path)

    for paragraph in doc.paragraphs:
        lines.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            line = '\t'.join([cell.text for cell in row.cells])
            lines.append(line)
        lines.append('')

    for rel in doc.part._rels:
        if "image" in doc.part._rels[rel].target_ref:
            lines.append(f"[Image: {doc.part._rels[rel].target_ref}]")
    
    for section in doc.sections:
        header = section.header
        footer = section.footer

        lines.append("=== Header ===")
        for paragraph in header.paragraphs:
            lines.append(paragraph.text)

        lines.append("=== Footer ===")
        for paragraph in footer.paragraphs:
            lines.append(paragraph.text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    return output_path
